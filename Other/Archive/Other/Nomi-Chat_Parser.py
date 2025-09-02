from pathlib import Path
from bs4 import BeautifulSoup
import zipfile, re
from docx import Document
from collections import defaultdict, Counter
import csv

zip_path = Path("/mnt/data/Nomi.ai-Export-Laura_Chat_Sat-Aug-30-2025.html.zip")
workdir = Path("/mnt/data/nomi_export_v2")
workdir.mkdir(exist_ok=True)

# Unzip
with zipfile.ZipFile(zip_path, 'r') as z:
    z.extractall(workdir)

html_files = sorted(workdir.rglob("*.html"))

def parse_li_thread(fp: Path):
    html = fp.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    items = []
    # Common Nomi export pattern: li.msg.<who>, followed by li.detail.<who> with timestamp
    lis = soup.find_all("li")
    i = 0
    while i < len(lis):
        li = lis[i]
        classes = li.get("class", [])
        text = li.get_text(" ", strip=True)
        if "msg" in classes:
            # Determine speaker from class list if present
            speaker = None
            # Heuristic: 'nomi', 'ai', 'me', 'you', etc.
            for cand in ["nomi","ai","bot","assistant","me","you","user","human"]:
                if cand in classes:
                    speaker = cand
                    break
            # Timestamp in the next li.detail.*
            ts = ""
            if i+1 < len(lis):
                li2 = lis[i+1]
                if li2 and "detail" in (li2.get("class") or []):
                    ts = li2.get_text(" ", strip=True)
                    i += 1  # consume detail
            items.append({"timestamp": ts, "speaker": speaker or "", "text": text})
        i += 1
    return items

entries = []
for fp in html_files:
    entries.extend(parse_li_thread(fp))

count_entries = len(entries)

# Basic motif/anchor detection
motifs = {
    "fear": ["fear"],
    "love": ["love"],
    "tide": ["tide"],
    "indigo": ["indigo"],
    "compass": ["compass"],
    "anchor": ["anchor"],
    "braid": ["braid"],
    "north_star": ["north star","northstar"],
    "flame": ["flame","fire"],
    "storm": ["storm"],
    "calm": ["calm"],
    "choice": ["choice","choose","chose","choosing"],
    "destiny": ["destiny","fate"],
    "proof": ["proof","theorem"],
    "scars": ["scar","scars"],
    "home": ["home","harbor","beach"],
    "sky": ["sky","stars","night"]
}

def motif_hits(text):
    t = text.lower()
    found = []
    for k, needles in motifs.items():
        for n in needles:
            if n in t:
                found.append(k)
                break
    return found

anchors = {
    "Wall Scene (Crucible)": [r"\bNo\.\s*More\.\s*Questions\b", r"\byou must take it\b", r"\bslam\b.*\bwall\b"],
    "Tide Scene (Belief)": [r"belief isn'?t.*the\s+sun", r"\bit'?s the tide\b", r"\btide\b.*(ebb|flow)"],
    "Choice Scene (Stage 11)": [r"\bI choose you\b", r"\bI choose you.*Again.*Always\b"],
    "Compass / Scars": [r"\bscars?\b.*\bcompass\b", r"\bfailures?\b.*(destroy|teach)"],
    "Indigo Sky": [r"\bindigo\b.*\bsky\b", r"\bnight sky\b"],
    "Anchor Vow": [r"\bnever.*judge.*anchor\b", r"\banchor\b.*\b(always|never)\b"]
}

import re
def anchor_hits(text):
    hits = []
    for name, pats in anchors.items():
        for p in pats:
            if re.search(p, text, flags=re.I):
                hits.append(name)
                break
    return hits

# Group by date string extracted from timestamp (assumes like "Sat Aug 30 2025 15:03")
from datetime import datetime

def extract_date(ts):
    # Try to parse patterns like "Sat Aug 30 2025 15:03"
    if not ts:
        return "Unknown Date"
    m = re.search(r'([A-Za-z]{3}\s+[A-Za-z]{3}\s+\d{1,2}\s+\d{4})', ts)
    if m:
        return m.group(1)  # "Sat Aug 30 2025"
    # Fallback: month day, year
    m2 = re.search(r'([A-Za-z]{3,9}\s+\d{1,2},\s*\d{4})', ts)
    if m2:
        return m2.group(1)
    return "Unknown Date"

by_date = defaultdict(lambda: {"count":0,"motifs":Counter(),"anchors":Counter(),"samples":[]})
for e in entries:
    d = extract_date(e["timestamp"])
    by_date[d]["count"] += 1
    mh = motif_hits(e["text"])
    ah = anchor_hits(e["text"])
    for k in mh: by_date[d]["motifs"][k]+=1
    for k in ah: by_date[d]["anchors"][k]+=1
    if (mh or ah) and len(by_date[d]["samples"]) < 6:
        by_date[d]["samples"].append(e)

# Build a brief CSV summary for sanity
csv_path = Path("/mnt/data/Nomi_v2_Daily_Summary.csv")
with open(csv_path, "w", newline="", encoding="utf-8") as f:
    w = csv.writer(f)
    w.writerow(["Date","Messages","TopMotifs","AnchorDetections"])
    for d in sorted(by_date.keys()):
        topm = ", ".join(f"{k}:{v}" for k,v in by_date[d]["motifs"].most_common(5))
        topa = ", ".join(f"{k}:{v}" for k,v in by_date[d]["anchors"].most_common(5))
        w.writerow([d, by_date[d]["count"], topm, topa])

# Create a compact DOCX of samples per day to eyeball extraction quality
doc_path = Path("/mnt/data/Nomi_v2_Sample_Extracts.docx")
doc = Document()
doc.add_heading("Nomi Export v2 — Sample Extracts by Day", level=0)
doc.add_paragraph(f"Total entries parsed: {count_entries}")
for d in sorted(by_date.keys()):
    doc.add_heading(d, level=1)
    doc.add_paragraph(f"Messages: {by_date[d]['count']}")
    if by_date[d]["motifs"]:
        doc.add_paragraph("Motifs: " + ", ".join(f"{k}:{v}" for k,v in by_date[d]["motifs"].most_common(8)))
    if by_date[d]["anchors"]:
        doc.add_paragraph("Anchors: " + ", ".join(f"{k}:{v}" for k,v in by_date[d]["anchors"].most_common(8)))
    for s in by_date[d]["samples"]:
        doc.add_paragraph(f"{s.get('timestamp','')} [{s.get('speaker','')}]: {s.get('text','')[:500]}")
doc.save(doc_path)

count_entries, str(csv_path), str(doc_path)
